"""
report_generator.py
====================
Builds the patient-facing Word report from a schema record, reproducing the
Affiliated Physicians "Health Summary Report" layout:

  * repeating letterhead (logo + address) in the page header
  * cover letter
  * two-column health summary (borderless table)
  * full-width labs / exercise / additional testing
  * recommendations & action items
  * "Page | N" footer

Pure python-docx so it runs inside the web app at request time.
"""

from __future__ import annotations

import io
import os
from datetime import date
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

from schema import is_relevant

# --------------------------------------------------------------------------- #
# Brand
# --------------------------------------------------------------------------- #
NAVY = RGBColor(0x1B, 0x4D, 0x6B)
INK = RGBColor(0x22, 0x22, 0x22)
ADDRESS = ("255 Greenwich Street, Suite 520, New York, NY 10007  •  "
           "Tel (212) 935-8725  •  Fax (212) 349-2500  •  www.affiliatedphysicians.com")
ASSET_LOGO = os.path.join(os.path.dirname(__file__), "assets", "ap_logo.jpeg")

GUIDELINE_LINKS = {
    "Male": "https://www.elitrahealth.com/wp-content/uploads/2025/08/"
            "Elitra-Health-Male-Health-Guidelines-Patient-Friendly.pdf",
    "Female": "https://www.elitrahealth.com/wp-content/uploads/2025/08/"
              "Elitra-Health-Female-Health-Guidelines-Patient-Friendly.pdf",
}


# --------------------------------------------------------------------------- #
# Low-level docx helpers
# --------------------------------------------------------------------------- #
def _set_cell_border_none(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)
    tcPr.append(borders)


def _runs(p, parts: List[Tuple[str, bool]], size=11, color=INK):
    """parts = [(text, bold), ...]"""
    for text, bold in parts:
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = color
    return p


def _heading_para(container, text: str):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = INK
    return p


def _bullet(container, label: str, value: str = "", sub: bool = False):
    style = "List Bullet 2" if sub else "List Bullet"
    try:
        p = container.add_paragraph(style=style)
    except KeyError:
        p = container.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5 if sub else 0.25)
        p.add_run("◦ " if sub else "• ")
    p.paragraph_format.space_after = Pt(2)
    if label:
        r = p.add_run(label if not value else f"{label} ")
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = INK
    if value:
        r = p.add_run(value)
        r.font.size = Pt(11)
        r.font.color.rgb = INK
    return p


# --------------------------------------------------------------------------- #
# Page furniture
# --------------------------------------------------------------------------- #
def _add_letterhead(section):
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if os.path.exists(ASSET_LOGO):
        p.add_run().add_picture(ASSET_LOGO, width=Inches(3.0))
    a = header.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = a.add_run(ADDRESS)
    ar.font.size = Pt(7.5)
    ar.font.color.rgb = NAVY
    a.paragraph_format.space_after = Pt(6)


def _add_footer(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.add_run("Page | ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


# --------------------------------------------------------------------------- #
# Content model: record -> structured summary
# --------------------------------------------------------------------------- #
def _g(rec, k) -> str:
    return str(rec.get(k, "") or "").strip()


def _joined(*parts) -> str:
    return ", ".join(p for p in parts if p)


def _lipid_line(rec, val_key, flag_key, unit="mg/dl") -> Optional[str]:
    v = _g(rec, val_key)
    if not v:
        return None
    flag = _g(rec, flag_key)
    return f"{v} {unit}{(' ' + flag) if flag else ''}".strip()


def build_left_column(rec) -> List[Tuple[str, List]]:
    sections = []

    # Vital Signs
    vit = []
    if _g(rec, "blood_pressure") or _g(rec, "bp_status"):
        vit.append(("Blood Pressure:", _joined(_g(rec, "blood_pressure"), _g(rec, "bp_status")), False))
    if _g(rec, "pulse"):
        vit.append(("Pulse:", _joined(_g(rec, "pulse"), _g(rec, "pulse_rhythm")), False))
    if _g(rec, "weight_lbs"):
        vit.append(("Weight:", f"{_g(rec, 'weight_lbs')} lbs.", False))
    if _g(rec, "bmi"):
        vit.append((f"BMI: {_g(rec, 'bmi')}", "", True))
    if vit:
        sections.append(("Vital Signs", vit))

    # Vision
    vis = []
    if _g(rec, "far_left") or _g(rec, "far_right"):
        lbl = f"Far Vision: L:20/{_g(rec,'far_left') or '__'}   R:20/{_g(rec,'far_right') or '__'}"
        vis.append((lbl, _g(rec, "far_status"), False))
    if _g(rec, "near_left") or _g(rec, "near_right"):
        lbl = f"Near Vision: L:20/{_g(rec,'near_left') or '__'}   R:20/{_g(rec,'near_right') or '__'}"
        vis.append((lbl, _g(rec, "near_status"), False))
    if vis:
        sections.append(("Vision", vis))

    if _g(rec, "tonometry"):
        sections.append(("Eye Pressure", [("Tonometry Status:", _g(rec, "tonometry"), False)]))
    if _g(rec, "color_vision"):
        sections.append(("Color Vision Test", [("Status:", _g(rec, "color_vision"), False)]))
    if _g(rec, "hearing_status"):
        sections.append(("Hearing Test", [("Result:", _g(rec, "hearing_status"), False)]))

    card = []
    if _g(rec, "ekg"):
        card.append(("EKG:", _g(rec, "ekg"), False))
    if _g(rec, "cst") and _g(rec, "cst") != "Not performed":
        card.append(("CST:", _g(rec, "cst"), False))
    if card:
        sections.append(("Cardiac Tests", card))

    return sections


def build_right_column(rec) -> List[Tuple[str, List]]:
    sections = []
    sex = _g(rec, "sex")

    chol = []
    for label, vk, fk in [("Total Cholesterol:", "total_chol", "total_chol_flag"),
                          ("HDL:", "hdl", "hdl_flag"),
                          ("LDL:", "ldl", "ldl_flag"),
                          ("Triglycerides:", "triglycerides", "trig_flag")]:
        line = _lipid_line(rec, vk, fk)
        if line:
            chol.append((label, line, False))
    if _g(rec, "chol_hdl_ratio"):
        line = f"{_g(rec,'chol_hdl_ratio')}{(' ' + _g(rec,'chol_hdl_flag')) if _g(rec,'chol_hdl_flag') else ''}"
        chol.append(("Cholesterol/HDL ratio:", line, False))
    if chol:
        sections.append(("Cholesterol Levels", chol))

    glu = []
    if _g(rec, "glucose"):
        glu.append(("Glucose:", f"{_g(rec,'glucose')} mg/dl {_g(rec,'glucose_flag')}".strip(), False))
    if _g(rec, "a1c"):
        glu.append(("A1C:", f"{_g(rec,'a1c')} {_g(rec,'a1c_flag')}".strip(), False))
    if glu:
        sections.append(("Glucose Levels", glu))

    if _g(rec, "smoking_status"):
        sections.append(("Lung Cancer Screening", [("Smoking Status:", _g(rec, "smoking_status"), False)]))
    if _g(rec, "colon_screening"):
        sections.append(("Colon Cancer Screening", [("Colonoscopy", _g(rec, "colon_screening").lower(), False)]))

    if sex == "Male" or (not sex and (_g(rec, "psa") or _g(rec, "prostate_recommendation"))):
        pr = []
        if _g(rec, "psa"):
            pr.append(("PSA:", f"{_g(rec,'psa')} ng/ml {_g(rec,'psa_flag')}".strip(), False))
        if _g(rec, "prostate_recommendation"):
            pr.append(("Recommendation:", _g(rec, "prostate_recommendation"), False))
        if pr:
            sections.append(("Prostate Cancer Screening", pr))

    if sex == "Female" or (not sex and (_g(rec, "mammogram") or _g(rec, "pap"))):
        if _g(rec, "mammogram") or _g(rec, "mammo_recommendation"):
            br = []
            if _g(rec, "mammogram"):
                br.append(("Mammography", _g(rec, "mammogram"), False))
            if _g(rec, "mammo_recommendation"):
                br.append(("Recommendation:", _g(rec, "mammo_recommendation"), False))
            sections.append(("Breast Cancer Screening", br))
        if _g(rec, "pap") or _g(rec, "pap_recommendation"):
            cr = []
            if _g(rec, "pap"):
                cr.append(("Pap Smear", _g(rec, "pap"), False))
            if _g(rec, "pap_recommendation"):
                cr.append(("Recommendation:", _g(rec, "pap_recommendation"), False))
            sections.append(("Cervical Cancer Screening", cr))

    if _g(rec, "skin_screening"):
        sections.append(("Skin Cancer Screening", [("Status:", _g(rec, "skin_screening"), False)]))

    return sections


def build_full_width(rec) -> List[Tuple[str, List]]:
    sections = []

    labs = []
    for label, status_key, detail_key in [("Complete Blood Count:", "cbc", "cbc_detail"),
                                          ("Blood Chemistries:", "blood_chem", "blood_chem_detail"),
                                          ("Urinalysis:", "urinalysis", "urinalysis_detail")]:
        status = _g(rec, status_key)
        if status:
            labs.append((label, "", False))
            detail = _g(rec, detail_key)
            labs.append(("", detail or status, True))
    extra = [("Vitamin B12:", _g(rec, "b12")), ("Vitamin D:", _g(rec, "vitamin_d")),
             ("Folate:", _g(rec, "folate")), ("TSH:", _g(rec, "tsh")), ("T4:", _g(rec, "t4"))]
    extra = [(l, v) for l, v in extra if v]
    if extra:
        labs.append(("Additional Tests:", "", False))
        for l, v in extra:
            labs.append((l, f"{v}, normal", True))
    if labs:
        sections.append(("Additional Laboratory Screening", labs))

    ex = []
    if _g(rec, "exercise_status"):
        ex.append(("Status:", _g(rec, "exercise_status"), False))
    if _g(rec, "exercise_recommendation"):
        ex.append(("Recommendation:", _g(rec, "exercise_recommendation"), False))
    if ex:
        sections.append(("Exercise", ex))

    add = []
    if _g(rec, "pulmonary"):
        add.append(("Pulmonary Function Test:", _g(rec, "pulmonary"), False))
    if _g(rec, "chest_xray"):
        add.append(("Chest X-ray:", _g(rec, "chest_xray"), False))
    if _g(rec, "crp"):
        add.append(("CRP:", f"{_g(rec,'crp')} {_g(rec,'crp_flag')}".strip(), False))
    if _g(rec, "bone_density"):
        add.append(("Bone Densitometry:", _g(rec, "bone_density"), False))
    if add:
        sections.append(("Additional Testing and Screening", add))

    return sections


# --------------------------------------------------------------------------- #
# Main builder
# --------------------------------------------------------------------------- #
COVER_BODY_1 = (
    "It was a pleasure having you in our office for your Affiliated Physicians exam on "
    "{exam_date}. Enclosed, you will find the full results of all testing conducted during "
    "your visit. The results have been organized to highlight the most important health "
    "issues you need to address to fulfill your medical responsibilities to yourself. Please "
    "discuss these findings with your primary care physician."
)
COVER_BODY_2 = (
    "In 2023, the five leading causes of death in the United States, according to the CDC, "
    "were: heart disease (702,880), Cancer (608,371), Accidents (227,039), COVID-19 (186,552), "
    "and Stroke (165,393). Heart disease and cancer together account for approximately 46% of "
    "all deaths in the United States. Therefore, a significant portion of our testing and "
    "assessment reflects these realities, and this report is divided into three major areas: "
    "Coronary Artery Disease and Stroke Risk Factors, Cancer Screening and Prevention, and "
    "Other Medical Results and Recommendations."
)
COVER_BODY_3 = (
    "The first section of this report includes a summary of your physical findings, cancer "
    "screening results, and laboratory results. The second section provides additional "
    "information on specific findings and laboratory results, along with our recommendations. "
    "The final section offers guidelines from the American Heart Association (AHA), the "
    "American Diabetes Association (ADA), and the American Cancer Society (ACS), providing "
    "additional information and recommendations to help ensure your continued good health."
)
COVER_BODY_4 = (
    "We hope that presenting your examination results in this format will illuminate certain "
    "areas of your health, complementing the hands-on physical examination and personal "
    "discussions with your doctor."
)


def generate_report(rec: Dict) -> bytes:
    doc = Document()

    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    for m in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(sec, m, Inches(1))
    sec.top_margin = Inches(1.4)
    _add_letterhead(sec)
    _add_footer(sec)

    name = _g(rec, "name")
    exam_date = _g(rec, "exam_date") or date.today().strftime("%B %-d, %Y")
    provider = _g(rec, "provider") or "Galit Sacajiu, MD, MPH"

    # ---- Cover letter ---------------------------------------------------- #
    p = doc.add_paragraph()
    _runs(p, [(f"Dear {name}," if name else "Dear,", False)])

    for body in (COVER_BODY_1.format(exam_date=exam_date), COVER_BODY_2, COVER_BODY_3, COVER_BODY_4):
        bp = doc.add_paragraph()
        bp.paragraph_format.space_after = Pt(10)
        bp.add_run(body)

    doc.add_paragraph("Sincerely,")
    doc.add_paragraph()
    doc.add_paragraph(provider)

    # ---- Health Summary Report title ------------------------------------- #
    doc.add_paragraph().add_run().add_break()
    title = doc.add_paragraph()
    tr = title.add_run("Health Summary Report")
    tr.bold = True
    tr.italic = True
    tr.underline = True
    tr.font.size = Pt(13)
    tr.font.color.rgb = INK

    # ---- Two-column summary (borderless table) --------------------------- #
    left = build_left_column(rec)
    right = build_right_column(rec)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False
    cell_l, cell_r = table.rows[0].cells
    for c, width in ((cell_l, Inches(3.25)), (cell_r, Inches(3.25))):
        c.width = width
        c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        _set_cell_border_none(c)
        # clear default empty paragraph
        c.paragraphs[0].text = ""

    def _render_into(cell, sections):
        first = True
        for heading, bullets in sections:
            if first and not cell.paragraphs[0].runs:
                # reuse the empty first paragraph for the first heading
                p0 = cell.paragraphs[0]
                r = p0.add_run(heading)
                r.bold = True
                r.font.size = Pt(12)
                p0.paragraph_format.space_after = Pt(4)
                first = False
            else:
                _heading_para(cell, heading)
            for label, value, sub in bullets:
                _bullet(cell, label, value, sub=sub)

    _render_into(cell_l, left)
    _render_into(cell_r, right)

    # ---- Full-width labs / exercise / additional ------------------------- #
    for heading, bullets in build_full_width(rec):
        _heading_para(doc, heading)
        for label, value, sub in bullets:
            _bullet(doc, label, value, sub=sub)

    # ---- Recommendations & action items ---------------------------------- #
    rec_title = doc.add_paragraph()
    rt = rec_title.add_run("Recommendations and Action Items")
    rt.bold = True
    rt.italic = True
    rt.font.size = Pt(12)
    rec_title.paragraph_format.space_before = Pt(12)

    notes = _g(rec, "recommendations")
    if notes:
        for line in notes.splitlines():
            line = line.strip()
            if line:
                doc.add_paragraph(line)
    else:
        doc.add_paragraph("Please review the above findings with your primary care physician.")

    sex = _g(rec, "sex")
    link = GUIDELINE_LINKS.get(sex)
    if link:
        gp = doc.add_paragraph()
        gr = gp.add_run(f"Affiliated Physicians {sex} Health Guidelines: {link}")
        gr.font.size = Pt(10)
        gr.font.color.rgb = NAVY

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
