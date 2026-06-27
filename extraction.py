"""
extraction.py
=============
Deterministic macro-sheet extraction. No LLM, no Vision API.

Pipeline:  raw bytes ──▶ plain text ──▶ field parsing ──▶ schema dict

Text is recovered with `pdfplumber` (PDF), `python-docx` (DOCX) or read
directly (TXT).  Parsing is two complementary strategies:

  1. **Checkbox tokenisation** – a filled form encodes a selection as a
     non-empty bracket, e.g. ``[x] Normal``.  `parse_checkbox_line`
     turns a line into (checked?, label) pairs which we map onto the
     known options for a field.

  2. **Label-anchored regex** – numeric blanks ("Total Cholesterol 187")
     are captured by a regex anchored on the field label.

Every value the parser is confident about is recorded in
``record["_meta"]["extracted_fields"]`` so the UI can show what came from
the document vs. what still needs a human.
"""

from __future__ import annotations

import io
import re
from typing import Dict, List, Optional, Tuple

from schema import blank_record, SECTIONS


# --------------------------------------------------------------------------- #
# 1. Text recovery
# --------------------------------------------------------------------------- #
def bytes_to_text(data: bytes, filename: str) -> str:
    """Recover plain text from a PDF / DOCX / TXT upload."""
    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return _pdf_to_text(data)
    if name.endswith(".docx"):
        return _docx_to_text(data)
    # txt / md / anything else: decode best-effort
    try:
        return data.decode("utf-8", errors="replace")
    except Exception:
        return data.decode("latin-1", errors="replace")


def _pdf_to_text(data: bytes) -> str:
    import pdfplumber

    chunks: List[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for page in pdf.pages:
            txt = page.extract_text(x_tolerance=1.5, y_tolerance=3) or ""
            chunks.append(txt)
    return "\n".join(chunks)


def _docx_to_text(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    lines: List[str] = []

    def _para_text(p) -> str:
        # Preserve a marker for checked boxes even if styling is lost.
        return p.text

    for block in doc.paragraphs:
        lines.append(_para_text(block))
    for table in doc.tables:
        for row in table.rows:
            lines.append(" ".join(c.text for c in row.cells))
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# 2. Low-level helpers
# --------------------------------------------------------------------------- #
# A checkbox is "checked" when something other than whitespace sits between
# the brackets: [x] [X] [✓] [√] [•] [*] ...  An empty [ ] or [] is unchecked.
_BOX = re.compile(r"\[\s*([^\]]*?)\s*\]")
_CHECK_CHARS = set("xX✓✔√•◆●*✗✘☑Xv")


def _is_checked(inner: str) -> bool:
    inner = (inner or "").strip()
    if not inner:
        return False
    # treat any non-space glyph as a mark
    return any((not ch.isspace()) for ch in inner)


def parse_checkbox_line(line: str) -> List[Tuple[bool, str]]:
    """
    Split a line into (checked, label) pairs.

    "[x] Normal [ ] Elevated"  ->  [(True, "Normal"), (False, "Elevated")]
    The label is the text after a box up to the next box.
    """
    pairs: List[Tuple[bool, str]] = []
    matches = list(_BOX.finditer(line))
    for i, m in enumerate(matches):
        checked = _is_checked(m.group(1))
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(line)
        label = line[start:end].strip(" .:\t")
        pairs.append((checked, label))
    return pairs


def _norm(s: str) -> str:
    return re.sub(r"[^a-z0-9]", "", (s or "").lower())


def _match_option(label: str, options: List[str]) -> Optional[str]:
    """Map a free-text checkbox label onto the closest known schema option."""
    nl = _norm(label)
    if not nl:
        return None
    # exact / containment match on normalised strings
    best = None
    for opt in options:
        no = _norm(opt)
        if not no:
            continue
        if nl == no:
            return opt
        if no in nl or nl in no:
            # prefer the longest overlap
            if best is None or len(no) > len(_norm(best)):
                best = opt
    return best


def _choice_from_lines(lines: List[str], anchors: List[str],
                       options: List[str]) -> Optional[str]:
    """
    Find the first line containing any anchor, then return the option whose
    checkbox is checked on that line (searching a couple of following lines
    for wrapped forms).
    """
    for idx, line in enumerate(lines):
        if any(a.lower() in line.lower() for a in anchors):
            window = " ".join(lines[idx:idx + 2])
            for checked, label in parse_checkbox_line(window):
                if checked:
                    opt = _match_option(label, options)
                    if opt:
                        return opt
    return None


def _value_after(lines: List[str], pattern: str, group: int = 1) -> Optional[str]:
    rx = re.compile(pattern, re.IGNORECASE)
    for line in lines:
        m = rx.search(line)
        if m:
            val = (m.group(group) or "").strip(" .:\t")
            if val and val not in {"_", "__", "___", "____"} and "___" not in val:
                return val
    return None


def _lab_status(lines: List[str], anchor: str,
                abnormal_opts: List[str]) -> Tuple[Optional[str], Optional[str]]:
    """
    For CBC / Blood Chemistries / Urinalysis lines: return (status, detail).
    'Normal' box ticked -> ("Normal", None). Otherwise collect ticked abnormal
    findings into a detail string and mark status 'Abnormal - see below'.
    """
    for idx, line in enumerate(lines):
        if anchor.lower() not in line.lower():
            continue
        window = " ".join(lines[idx:idx + 2])
        ticked = [label for checked, label in parse_checkbox_line(window) if checked]
        if not ticked:
            return None, None
        norm_ticked = [_norm(t) for t in ticked]
        if any(_norm("Normal") == nt or _norm("Normal") in nt for nt in norm_ticked):
            return "Normal", None
        details = []
        for opt in abnormal_opts:
            for t in ticked:
                if _norm(opt) in _norm(t) or _norm(t) in _norm(opt):
                    details.append(opt)
                    break
        if details:
            return "Abnormal - see below", ", ".join(details)
        return "Abnormal - see below", None
    return None, None


def _trailing_block(text: str, anchors: List[str]) -> Optional[str]:
    """Return cleaned free text appearing after the last matching anchor line."""
    lines = text.splitlines()
    start = None
    for i, line in enumerate(lines):
        if any(a.lower() in line.lower() for a in anchors):
            start = i
    if start is None:
        return None
    # text on the anchor line after the colon, plus subsequent non-empty lines
    tail = []
    head = lines[start].split(":", 1)
    if len(head) == 2 and head[1].strip():
        tail.append(head[1].strip())
    for line in lines[start + 1:]:
        s = line.strip()
        if not s:
            continue
        if "MACRO" in s and len(s) < 12:
            continue
        tail.append(s)
    cleaned = [t for t in tail if t and t.lower() not in {"page | 2", "page | 3"}]
    return "\n".join(cleaned) if cleaned else None


# --------------------------------------------------------------------------- #
# 3. Field parsers
# --------------------------------------------------------------------------- #
def extract_record(text: str, source_file: str = "") -> Dict:
    rec = blank_record()
    rec["_meta"]["source_file"] = source_file
    found: List[str] = []
    raw_lines = [re.sub(r"[ \t]+", " ", ln).strip() for ln in text.splitlines()]
    lines = [ln for ln in raw_lines if ln]

    def setv(key: str, value: Optional[str]):
        if value not in (None, ""):
            rec[key] = value
            found.append(key)

    # --- Patient / visit --------------------------------------------------- #
    setv("name", _value_after(lines, r"Patient Name[:\s]+([A-Za-z][A-Za-z .,'-]+?)(?:\s+ID|\s{2,}|$)"))
    setv("patient_id", _value_after(lines, r"\bID[:\s_]*([A-Za-z0-9-]{2,})"))
    setv("exam_date", _value_after(lines, r"Exam Date[:\s]+([A-Za-z0-9 ,/-]+?)$"))
    # sex from explicit header line MALE / FEMALE
    head = " ".join(lines[:4]).upper()
    if re.search(r"\bFEMALE\b", head):
        setv("sex", "Female")
    elif re.search(r"\bMALE\b", head):
        setv("sex", "Male")

    # --- Vitals ------------------------------------------------------------ #
    setv("blood_pressure", _value_after(lines, r"Blood Pressure[:\s]*\[?\s*\]?\s*([0-9]{2,3}\s*/\s*[0-9]{2,3})"))
    setv("bp_status", _choice_from_lines(lines, ["Blood Pressure"],
                                         ["Normal", "Borderline Elevated", "Elevated", "Very high"]))
    setv("pulse", _value_after(lines, r"Pulse[:\s]*([0-9]{2,3})\b"))
    rhythm = _choice_from_lines(lines, ["Pulse"], ["Reg", "Irreg"])
    if rhythm == "Reg":
        rhythm = "Regular"
    elif rhythm == "Irreg":
        rhythm = "Irregular"
    setv("pulse_rhythm", rhythm)
    setv("pulse_ox", _value_after(lines, r"Pulse Ox%?[:\s]*([0-9]{1,3})"))
    setv("temperature", _value_after(lines, r"Temperature[:\s_]*([0-9]{2,3}\.?[0-9]?)"))
    setv("height_in", _value_after(lines, r"Height[:\s_]*([0-9]{2,3})\s*in"))
    setv("weight_lbs", _value_after(lines, r"Weight[:\s_]*([0-9]{2,3})\s*lbs"))
    setv("bmi", _value_after(lines, r"BMI[:\s_]*([0-9]{1,2}\.?[0-9]?)"))
    if _choice_from_lines(lines[:2], ["medication"], ["on medication"]):
        setv("on_bp_medication", "Yes")

    # --- Vision ------------------------------------------------------------ #
    setv("far_left", _value_after(lines, r"Far:?\s*Left\s*20\s*/\s*([0-9]{2,3})"))
    setv("far_right", _value_after(lines, r"Far:.*?Right\s*20\s*/\s*([0-9]{2,3})"))
    setv("near_left", _value_after(lines, r"Near:?\s*Left\s*20\s*/\s*([0-9]{2,3})"))
    setv("near_right", _value_after(lines, r"Near:.*?Right\s*20\s*/\s*([0-9]{2,3})"))
    setv("far_status", _choice_from_lines(lines, ["Far: ["],
                                          ["Normal", "Decreased right", "Decreased left", "Decreased both"]))
    setv("near_status", _choice_from_lines(lines, ["Near: ["],
                                           ["Normal", "Decreased right", "Decreased left", "Decreased both"]))
    setv("correction", _choice_from_lines(lines, ["correction"], ["with", "without"]))
    setv("tonometry", _choice_from_lines(lines, ["Tonometry"], ["Normal", "High"]))
    setv("color_vision", _choice_from_lines(lines, ["Color Vision"], ["Normal", "Abnormal", "See below"]))

    # --- Hearing ----------------------------------------------------------- #
    hearing = _choice_from_lines(
        lines, ["Decr R", "Chronic hearing", "Audiologist", "Hearing"],
        ["Normal", "Decr R", "Decr L", "Decr HF R", "Decr HF L",
         "Chronic hearing loss", "Bilateral"])
    hearing_map = {
        "Decr R": "Decreased right", "Decr L": "Decreased left",
        "Decr HF R": "High frequency hearing loss right",
        "Decr HF L": "High frequency hearing loss left",
        "Chronic hearing loss": "Chronic hearing loss", "Bilateral": "Bilateral loss",
    }
    setv("hearing_status", hearing_map.get(hearing, hearing))

    # --- Cardiac ----------------------------------------------------------- #
    setv("ekg", _choice_from_lines(lines, ["EKG"],
                                   ["Normal", "Abnormal, see below", "Normal, see below", "See below"]))
    setv("cst", _choice_from_lines(lines, ["CST"],
                                   ["Normal", "Abnormal, see below", "Normal, see below",
                                    "Negative for Ischemia", "See below"]))

    # --- Cholesterol ------------------------------------------------------- #
    for key, label in [("total_chol", "Total Cholesterol"), ("hdl", "HDL Cholesterol"),
                       ("ldl", "LDL Cholesterol"), ("triglycerides", "Triglycerides"),
                       ("chol_hdl_ratio", "Cholesterol/HDL")]:
        val = _value_after(lines, rf"{re.escape(label)}[:\s_]*([0-9]+\.?[0-9]?)")
        setv(key, val)
        # trailing flag word, e.g. "Total Cholesterol 187 Good"
        flag = _value_after(lines, rf"{re.escape(label)}[:\s_]*[0-9.]+\s+(Good|Normal|Borderline|Elevated|High|Low|All good)",)
        if flag:
            fk = {"total_chol": "total_chol_flag", "hdl": "hdl_flag", "ldl": "ldl_flag",
                  "triglycerides": "trig_flag", "chol_hdl_ratio": "chol_hdl_flag"}[key]
            setv(fk, flag.title() if flag.lower() != "all good" else "All Good")

    # --- Glucose ----------------------------------------------------------- #
    setv("glucose", _value_after(lines, r"Glucose[:\s_]*([0-9]{2,3})"))
    setv("glucose_flag", _choice_from_lines(lines, ["Glucose"], ["Normal", "Borderline", "Elevated"]))
    setv("a1c", _value_after(lines, r"A1c[:\s_]*([0-9]{1,2}\.?[0-9]?)"))
    setv("a1c_flag", _choice_from_lines(lines, ["A1c"], ["Normal", "Borderline", "Elevated", "Pre-diabetes"]))

    # --- Cancer screening -------------------------------------------------- #
    setv("smoking_status", _choice_from_lines(lines, ["Smoking"], ["Smoker", "Non-smoker", "Former smoker"]))
    setv("colon_screening", _choice_from_lines(lines, ["Colon cancer"],
                                               ["Up-to-date", "Recommended", "at 45 yo", "See below"]))
    # Prostate
    setv("prostate_exam", _choice_from_lines(lines, ["Prostate Exam"],
                                             ["Normal", "UTD", "Declined", "N/A", "Other"]))
    setv("psa", _value_after(lines, r"PSA[:\s_]*([0-9]+\.?[0-9]*)"))
    setv("psa_flag", _choice_from_lines(lines, ["PSA"], ["Normal", "Elevated", "See below"]))
    setv("prostate_recommendation", _choice_from_lines(
        lines, ["Annual prostate", "prostate screening"],
        ["Annual prostate screening", "UTD", "Recommended at 40", "Recommended at 45",
         "Recommended at 50", "Follow up with Urologist"]))
    # Breast / cervical / ovarian
    setv("mammogram", _choice_from_lines(lines, ["Mammogram ["],
                                         ["Normal", "Normal with fibrocystic tissue",
                                          "Abnormal, see attached report", "See below"]))
    setv("pap", _choice_from_lines(lines, ["Pap Test"],
                                   ["Normal", "Normal with HPV", "Abnormal, see below", "See below"]))
    setv("gyn_exam", _choice_from_lines(lines, ["Gynecological exam"],
                                        ["Up-to-date", "Recommended", "See below"]))
    setv("skin_screening", _choice_from_lines(lines, ["Dermatology"],
                                              ["UTD", "Family hx of skin cancer", "Recommended", "See below"]))

    # --- Labs -------------------------------------------------------------- #
    # CBC / Chemistries / Urinalysis: "Normal" if that box is ticked,
    # otherwise capture whichever abnormal finding is ticked as the detail.
    for anchor, status_key, detail_key, abnormal_opts in [
        ("Complete Blood Count", "cbc", "cbc_detail",
         ["WBC", "Platelets", "Anemia", "Hemoglobin", "Hematocrit"]),
        ("Blood Chemistries", "blood_chem", "blood_chem_detail",
         ["ALT", "AST", "GGT", "Iron", "Transferrin Saturation", "Bilirubin", "Uric Acid"]),
        ("Urinalysis", "urinalysis", "urinalysis_detail",
         ["WBC", "RBC", "Squamous Epithelial cells", "Bacteria",
          "Leukocyte Esterase", "Protein", "Ketones"]),
    ]:
        status, detail = _lab_status(lines, anchor, abnormal_opts)
        if status:
            setv(status_key, status)
        if detail:
            setv(detail_key, detail)
    setv("tsh", _value_after(lines, r"\bTSH\b[:\s]*([0-9]+\.?[0-9]*)"))
    setv("t4", _value_after(lines, r"\bT4\b[:\s]*([0-9]+\.?[0-9]*)"))
    setv("b12", _value_after(lines, r"B12[:\s]*([0-9]+\.?[0-9]*)"))
    setv("folate", _value_after(lines, r"Folate[:\s]*([0-9]+\.?[0-9]*)"))
    setv("vitamin_d", _value_after(lines, r"Vitamin D[:\s]*([0-9]+\.?[0-9]*)"))

    # --- Other ------------------------------------------------------------- #
    setv("exercise_status", _choice_from_lines(lines, ["Exercise Status"],
                                               ["Regular exercise", "Limited exercise", "No exercise"]))
    setv("exercise_recommendation", _choice_from_lines(
        lines, ["Continue exercise", "Cardiovascular exercises", "Resistance exercises"],
        ["Continue exercise regimen", "Cardiovascular exercises recommended",
         "Resistance exercises recommended", "See below"]))
    setv("crp", _value_after(lines, r"CRP[:\s_]*([0-9]+\.?[0-9]*)"))
    setv("crp_flag", _choice_from_lines(lines, ["CRP"], ["Normal", "Elevated"]))
    setv("pulmonary", _choice_from_lines(lines, ["Pulmonary Function"],
                                         ["Normal", "Obstructive pattern", "Restrictive pattern",
                                          "Mild", "Moderate"]))
    setv("chest_xray", _choice_from_lines(lines, ["Chest X-Ray"], ["Normal", "Other"]))
    setv("bone_density", _choice_from_lines(lines, ["Bone Densitometry"],
                                            ["Normal", "Up-to-date", "Recommended", "Other"]))

    # Free-text recommendations: everything after "All Other Recommendations:"
    rec_note = _trailing_block(text, ["All Other Recommendations", "Additional Notations"])
    if rec_note:
        setv("recommendations", rec_note)

    rec["_meta"]["extracted_fields"] = sorted(set(found))
    return rec


# --------------------------------------------------------------------------- #
# 4. Optional clinical flag derivation (rule-based, fully transparent)
# --------------------------------------------------------------------------- #
def derive_flags(rec: Dict) -> Dict[str, str]:
    """
    Suggest standard reference-range flags for numeric labs that have a value
    but no flag yet. Returns {field: suggested_flag}. These are conventional
    adult thresholds, surfaced as *suggestions* for the reviewer — never
    applied silently.
    """
    out: Dict[str, str] = {}

    def num(key):
        try:
            return float(str(rec.get(key, "")).replace(",", ""))
        except (TypeError, ValueError):
            return None

    rules = {
        "total_chol_flag": (num("total_chol"), [(200, "Good"), (240, "Borderline"), (10**9, "Elevated")]),
        "ldl_flag": (num("ldl"), [(100, "Good"), (130, "Normal"), (160, "Borderline"), (10**9, "Elevated")]),
        "trig_flag": (num("triglycerides"), [(150, "Good"), (200, "Borderline"), (10**9, "Elevated")]),
        "glucose_flag": (num("glucose"), [(100, "Normal"), (126, "Borderline"), (10**9, "Elevated")]),
        "a1c_flag": (num("a1c"), [(5.7, "Normal"), (6.5, "Borderline"), (10**9, "Pre-diabetes")]),
    }
    for flag_key, (value, bands) in rules.items():
        if value is None or rec.get(flag_key):
            continue
        for ceiling, label in bands:
            if value < ceiling:
                out[flag_key] = label
                break

    # HDL: higher is better (>=60 Good, 40-59 Normal, <40 Low)
    hdl = num("hdl")
    if hdl is not None and not rec.get("hdl_flag"):
        out["hdl_flag"] = "Good" if hdl >= 60 else ("Normal" if hdl >= 40 else "Low")

    # Chol/HDL ratio: <3.5 All Good, <5 Normal, else Elevated
    ratio = num("chol_hdl_ratio")
    if ratio is not None and not rec.get("chol_hdl_flag"):
        out["chol_hdl_flag"] = "All Good" if ratio < 3.5 else ("Normal" if ratio < 5 else "Elevated")

    # BMI category note isn't a schema field; BP status from systolic/diastolic
    bp = str(rec.get("blood_pressure", ""))
    m = re.match(r"\s*(\d{2,3})\s*/\s*(\d{2,3})", bp)
    if m and not rec.get("bp_status"):
        sys, dia = int(m.group(1)), int(m.group(2))
        if sys < 120 and dia < 80:
            out["bp_status"] = "Normal"
        elif sys < 130 and dia < 80:
            out["bp_status"] = "Borderline Elevated"
        elif sys < 140 or dia < 90:
            out["bp_status"] = "Elevated"
        else:
            out["bp_status"] = "Very high"

    return out


def coverage(rec: Dict) -> Tuple[int, int]:
    """(#fields with a value, #total fields) ignoring meta/detail helper fields."""
    total = 0
    filled = 0
    skip = {"cbc_detail", "blood_chem_detail", "urinalysis_detail"}
    for section in SECTIONS:
        for key, *_ in section["fields"]:
            if key in skip:
                continue
            total += 1
            if str(rec.get(key, "")).strip():
                filled += 1
    return filled, total
